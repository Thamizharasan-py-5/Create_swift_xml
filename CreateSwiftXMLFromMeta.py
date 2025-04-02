import os
from zipfile import ZipFile
import zipfile36 as zipfile
from lxml import etree
import time
import imaplib, email
from datetime import date
import logging
import urllib.request
import subprocess
from lxml.etree import tostring
from itertools import chain
import ftplib
from ftplib import FTP
import docx
import re
import docx2txt
import shutil
import requests
import tkinter as tk
from tkinter import messagebox
import sys

createdSwiftXML='//j-fs01/OUP_Journals-L/OUP_File_Autodownload/'
configFile='//j-fs01/OUP_Journals-L/OUP_File_Autodownload/Config_file/Config.xml'
swiftxmlpath='//j-fs01/OUP_Journals-L/OUP_File_Autodownload/Config_file/JEXBOT-2022-307097.xml'
metaTxtFile='//j-fs01/OUP_Journals-L/OUP_File_Autodownload/Config_file/Swiftxml_creation.txt'
user = 'oupjbookings@newgen.co'
password = 'June@2024'
imap_url = 'imap.gmail.com'
imap_port= 993
months={}
months['Jan']='01'
months['Feb']='02'
months['Mar']='03'
months['Apr']='04'
months['May']='05'
months['Jun']='06'
months['Jul']='07'
months['Aug']='08'
months['Sep']='09'
months['Oct']='10'
months['Nov']='11'
months['Dec']='12'

months['January']='01'
months['February']='02'
months['March']='03'
months['April']='04'
months['May']='05'
months['June']='06'
months['July']='07'
months['August']='08'
months['September']='09'
months['October']='10'
months['November']='11'
months['December']='12'

days={}

days['1']='01'
days['2']='02'
days['3']='03'
days['4']='04'
days['5']='05'
days['6']='06'
days['7']='07'
days['8']='08'
days['9']='09'
days['10']='10'
days['11']='11'
days['12']='12'
days['13']='13'
days['14']='14'
days['15']='15'
days['16']='16'
days['17']='17'
days['18']='18'
days['19']='19'
days['20']='20'
days['21']='21'
days['22']='22'
days['23']='23'
days['24']='24'
days['25']='25'
days['26']='26'
days['27']='27'
days['28']='28'
days['29']='29'
days['30']='30'
days['31']='31'

con = imaplib.IMAP4_SSL(imap_url, imap_port)
con.login(user, password)
con.select('Inbox')
file = ''

def upload_xmlfile(xmlFileUpload):
    curl_executable = os.path.join(os.path.dirname(__file__), '//j-fs01/OUP_Journals-L/OUP_File_Autodownload/Issue_compilation/Config_file/curl-8.9.1_3-win64-mingw/bin/curl.exe')
    try:
        with open(xmlFileUpload, 'r') as f:
            xml_data = f.read()
    except FileNotFoundError:
        print("Error: The specified XML file was not found.")
        exit(1)
    curl_command = [
        curl_executable,
        '--location',
        '--request', 'POST',
        'https://api.oup.com/swift/v1/Manuscripts',
        '--header', 'Content-Type: application/xml',
        '--header', 'Ocp-Apim-Subscription-Key : bc118aad8ea342eba54caf6c351333e0',
        '--data', xml_data,
        '--ssl-no-revoke'
    ]
    try:
        result = subprocess.run(curl_command, capture_output=True, text=True, check=True)
        print("Status Code:", result.returncode)
        print("Response Text:", result.stdout)
        response_text=result.stdout
        clean_xml = response_text.split("?>", 1)[-1].strip()
        #if  in e.stderr:
        root = etree.fromstring(clean_xml) # Extract Messages element text 
        #messages = root.find("Messages")
        messages0= root.findall("Messages")
        message=xmlFileUpload.split('/')[-1]
        message=message.replace('.xml','')
        try:
            for m in messages0:
                message_text = m.text.strip() if messages0 is not None else "No message found"
                if 'Manuscripts inserted' in message_text:
                    add_data=message+' '+message_text+'\n'
                    with open(upl_file, "a") as file:
                        file.write(add_data+".\n")
                    print(f"Message: {message_text}")
                if 'A manuscript was found with the same number' in message_text:
                    title=f"This manuscript number : {message} is already upload\n"
                    with open(upl_file, "a") as file:
                        file.write(title+".\n")
        except Exception as ex12:
            print(ex12)
    except subprocess.CalledProcessError as e:
        print("Error Code:", e.returncode)
        print("Error Output:", e.stderr)
        
def upload_file(zipFileUpload):

    try:
        ftp_host = 'ftp.newgenimaging.com'
        ftp_user = 'oupjournals'
        ftp_pass = 'uEmg3n9upJ0UrNl'
        zipFileName = os.path.basename(zipFileUpload)
        remote_file_path = "/From_OUP/EXBOTJ/reprocess/"+zipFileName
        ftp = FTP(ftp_host)
        ftp.login(user=ftp_user, passwd=ftp_pass)
        print("✅ Connected and logged in to FTP server")
        with open(zipFileUpload, 'rb') as file:
            ftp.storbinary(f'STOR {remote_file_path}', file)
            print(f"✅ Uploaded to '{remote_file_path}")
        ftp.quit()
        print("✅ FTP connection closed")
    except:
        print ("unable to connect ftp server")

def stringify_children(node):
    
    parts = ([node.text] +
            list(chain(*([c.text, tostring(c), c.tail] for c in node.getchildren()))) +
            [node.tail])
    # filter removes possible Nones in texts and tails
    result = ''.join(map(lambda item: item if isinstance(
        item, str) else "", parts))
    return result

def readtxt(filename,xmlfileparser):
    authors=[]
    affliations={}
    authorAff={}
    foreNames=xmlfileparser.xpath('//Contacts/Contact/ForeName')
    surNames=xmlfileparser.xpath('//Contacts/Contact/SurName')

    for index,f in enumerate(foreNames):
        author=f.text+' '+surNames[index].text
        authors.append(author)
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        pTxt=p.text
        if(len(pTxt)<=1):
            continue
        firstCh=pTxt[0]
        if firstCh.isnumeric()==True:
            affliations[firstCh]=pTxt
            
    fullText = []
    for a in authors:
        for para in doc.paragraphs:
            paratxt=para.text
            if a.upper() in paratxt.upper():
                my_list=paratxt.split(', ')
                idx=0
                for l in my_list:
                    if a.lower() in l.lower():
                        if l[len(a)] in affliations:
                            
                            authorAff[a]=affliations[l[len(a)]]
                break


    for index,f in enumerate(foreNames):
        authorTxt1=f.text+' '+surNames[index].text
        if authorTxt1 not in authorAff:
            continue
        
        aff=authorAff[authorTxt1]
        aff=aff[1:]
        addressSplited=aff.split(', ')
        address=f.getparent().xpath('./Addresses/Address')
        firstAdd=address[0]
        childElts=firstAdd.getchildren()
        i=0
        postalCode=''
        postalCodeElt=None
        for c in childElts:
            try:
                addTxt=addressSplited[i]
                if addTxt.isnumeric()==True:
                    postalCode=addTxt
                if c.tag=='PostCode':
                    #postalCodeElt=c
                    continue
                if c.tag=='County':
                    c.text=''
                    continue
                if c.tag=='Country':
                    c.text=addressSplited[len(addressSplited)-1]
                    continue
                c.text=addTxt
                i=i+1
            except Exception as ex:
                print (ex)
        if postalCodeElt is not None:
            postalCodeElt.text=postalCode

def MainFunction(zipFile):
    
    def get_body(msg):
        if msg.is_multipart():
            return get_body(msg.get_payload(0))
        else:
            return msg.get_payload(None, True)
    def search(key, value, con,fileName):
        searchStr='('+key+' "'+value+'" SUBJECT "'+fileName+'")'
        #searchStr='(FROM "j.exp.bot@lancaster.ac.uk" SUBJECT "Manuscript JEXBOT/2022/306989 accepted for publication in JXB")'
        #result, data = con.search(None, key, value, 'X-GM-RAW', 'subject: Manuscript JEXBOT/2022/307803 accepted for publication in JXB')
        result, data = con.search(None, searchStr)
        return data
    def get_emails(result_bytes):
        msgs = []
        message=''
        try:
            for num in result_bytes[0].split():
                _, data = con.fetch(num, '(RFC822)')
                _, raw_email = data[0]
                email_message = email.message_from_bytes(raw_email)
                print (email_message)
                for part in email_message.walk():
                    message = part.get_payload(decode=True)
                    print (message)
                    if part.get_content_type()=="text/plain" or part.get_content_type()=="text/html":
                        message = part.get_payload(decode=True)
                        print("Message: \n", message.decode())
                        print("==========================================\n")
                        break
        except Exception as ex:
            print(ex)
        return message
    con = imaplib.IMAP4_SSL(imap_url)
    user = 'oupjbookings@newgen.co'
    password = 'June@2024'
    con.login(user,password)
    con.select('Inbox')
    xmlfile=''
    filename = zipFile#sys.argv[1] #'C:/Thamizh/metaxml/jexbot-305887.zip'
    path=os.path.dirname(filename)
    print (path)
    with ZipFile(filename,'r') as zip:
        zip.extractall(path)
        print('Done!')
    baseFileName=os.path.basename(filename)
    baseFileName=baseFileName.replace('.zip','')
    newPath=path+'/'+baseFileName
    word_filePath=newPath+'/'+'suppl_data'
    try:
        dir_list1= os.listdir(word_filePath)
    except Exception as ex5:
        print(ex5)
    try:  
        dir_list = os.listdir(newPath)
        for f in dir_list:
            if '.xml' in f:
                xmlfile=f
    except Exception as ex9:
        print(ex9)
    try:
        for f1 in dir_list1:
            if '.docx' in f1:
                wordFile=f1
            print(wordFile)
        doc_Path=word_filePath+'/'+wordFile
    except Exception as ex1:
        print(ex1)
    metaxmlpath=newPath+'/'+xmlfile
    todays_date = date.today()
    currentYear=int(todays_date.year)
    fileName=os.path.basename(filename)
    fileName=fileName.replace('.zip','')
    fileNameSplitted=fileName.split('-')
    fileName='Manuscript '+str(fileNameSplitted[0]).upper()+'-'+str(currentYear)+'-'+fileNameSplitted[1]
    metaarticleIDTxt=str(fileNameSplitted[0]).upper()+'-'+str(currentYear)+'-'+fileNameSplitted[1]
    fileName=fileName.replace('-','/')
    msgs = get_emails(search('FROM', 'j.exp.bot@lancaster.ac.uk', con,fileName))
    if(msgs==''):
        for d in path:
            prevYear=currentYear
            currentYear=str(int(currentYear)-1)
            fileName=fileName.replace('/'+str(prevYear)+'/','/'+str(currentYear)+'/')
            metaarticleIDTxt=str(fileNameSplitted[0]).upper()+'-'+str(currentYear)+'-'+fileNameSplitted[1]
            msgs = get_emails(search('FROM', 'j.exp.bot@lancaster.ac.uk', con,fileName))
            if msgs:
                break    
    xmlStr=''
    with open(swiftxmlpath,'r') as f:
        xmlStr=f.read()
        xmlStr=xmlStr.replace('xmlns="','xmlns:temp="')
        with open(swiftxmlpath,'w') as w:
            w.write(xmlStr)
            w.close()
        f.close()
    configParser = etree.parse(configFile)
    
    xmlStr=''
    try:
        with open(metaxmlpath,'r',encoding="utf-8") as f:
            xmlStr=f.read()
            xmlStr=xmlStr.replace('&','&amp;')
            if '<article xmlns:xlink="http://www.w3.org/1999/xlink" ' not in xmlStr:
                xmlStr=xmlStr.replace('<article ','<article xmlns:xlink="http://www.w3.org/1999/xlink" ')
            with open(metaxmlpath,'w',encoding="utf-8") as w:
                w.write(xmlStr)
                w.close()
                f.close()
    except Exception as eexx:
        print(eexx)
    metaxmlparser = etree.parse(metaxmlpath)
    
    swiftxmlparser = etree.parse(swiftxmlpath)
    
    elts=configParser.xpath('//element')
    
    articleIDs=metaxmlparser.xpath('//article-meta/article-id')
    if(len(articleIDs)>0):
        articleID=articleIDs[0]
        articleIDMSID=etree.Element('article-id')
        articleIDMSID.attrib['pub-id-type']='ms_no'
        articleIDMSID.text=metaarticleIDTxt
        articleID.addnext(articleIDMSID)
    metaxmlparser.write(metaxmlpath,encoding="utf-8")
    output = subprocess.run(["C:/Program Files/Adobe/Adobe InDesign 2020/Plug-Ins/AutoComp/UTF8.exe",metaxmlpath], capture_output=True)
    for e in elts:
        findXpath=e.attrib['find']
        replaceXpath=e.attrib['replace']
        newStr=''
        findText=metaxmlparser.xpath('string('+findXpath+')')
        append=''
        if 'append' in e.attrib:
            append=e.attrib['append']
        replaceElts=swiftxmlparser.xpath(replaceXpath)
        
        if(len(replaceElts)>0):
            splitVal=''
            posVal=''
            attVal=''
            if 'att' in e.attrib:
                attVal=e.attrib['att']
            if 'split' in e.attrib:
                if 'pos' in e.attrib:
                    splitVal=e.attrib['split']
                    posVal=e.attrib['pos']
                    
            if 'append' in e.attrib:
                newStr=str(replaceElts[0].text)
                if newStr=='None':
                    newStr=''
                findText=newStr+''+findText+''+append
            
            if splitVal!='' and posVal!='':
                findText=findText.split(splitVal)[int(posVal)]
            if 'att' not in e.attrib:
                replaceElts[0].text=findText
            else:
                replaceElts[0].attrib[attVal]=findText
    
    def get_body(msg):
        if msg.is_multipart():
            return get_body(msg.get_payload(0))
        else:
            return msg.get_payload(None, True)
    def search(key, value, con):
        searchStr1='('+key+' "'+value+'" SUBJECT "'+fileName+'")'
        #result, data = con.search(None, key, value, 'X-GM-RAW', 'subject: Manuscript JEXBOT/2022/307803 accepted for publication in JXB')
        result, data = con.search(None, searchStr1)
        
        return data
    
    
    msgs = get_emails(search('FROM', 'j.exp.bot@lancaster.ac.uk', con))
    try:
        docpath = metaTxtFile
        doc = open(docpath,'w',encoding='utf-8')
        doc.write(msgs.decode())
        doc.close()
    except Exception as ex1:
        print(ex1)
    notes=''
    isNote=False
    infoMap={}
    noOfFigures=0
    noOfColorFigures=0
    with open(docpath,'r',encoding="utf-8") as f:
         file = f.readlines()
         for r in file:
             wordArr=r.split(':')
             firstWord=wordArr[0]
             if('Number of Figures' in firstWord):
                 noOfFigures=int(wordArr[1].strip())
             if('Number of Colour Figures' in firstWord):
                 noOfColorFiguresStr=wordArr[1].strip()
                 if noOfColorFiguresStr.isnumeric()==True:
                     noOfColorFigures=int(noOfColorFiguresStr)
             
             if ('Revision Received Date'==firstWord or'Original Received Date'==firstWord or'Number of Tables'==firstWord or'Institution'==firstWord or'Email'==firstWord or'Fax'==firstWord or'Phone'==firstWord or'Zip/Postal Code'==firstWord or'Title'==firstWord or'Authors'==firstWord or 'Corresponding Author'==firstWord or 'Accepted Date'==firstWord or 'Manuscript Number'==firstWord or 'Address'==firstWord or 'Note  1.'):
                 infor=r.replace(firstWord+':','',1)
                 if(firstWord not in infoMap):
                     infoMap[firstWord]=infor
             if(isNote==True):
                 if('----' not in firstWord and '\n'!=firstWord):
                     notes=notes+''+r
             if(firstWord=='Original Received Date'):
                 notes='\n'+notes+''+r
                 isNote=True
             
             if('ISSUE' in r):
                 isNote=False
    figuresElts=swiftxmlparser.xpath('//Figures')
    noOfFiguresElts=swiftxmlparser.xpath('//Figures/Figure')
    if(len(noOfFiguresElts)>noOfFigures):
        balCnt=len(noOfFiguresElts)-noOfFigures
        cnt=0
        for fig in reversed(noOfFiguresElts):
            if(balCnt<=cnt):
                break
            else:
                fig.getparent().remove(fig)
                cnt=cnt+1
    elif (len(noOfFiguresElts)<noOfFigures):
        balCnt=noOfFigures-len(noOfFiguresElts)
        for c in range(balCnt):
            figElt=etree.Element('Figure')
            colorPrint=etree.Element('ColourPrint')
            colorOnline=etree.Element('ColourOnline')
            figuresElts[0].append(figElt)
            figElt.append(colorPrint)
            figElt.append(colorOnline)
    
    manuScriptNos=swiftxmlparser.xpath('//Manuscript/Number')
    manuScriptNos[0].text=metaarticleIDTxt
    tableElts=swiftxmlparser.xpath('//TableCount')
    if 'Number of Tables' in infoMap:
        tableEltsTxt=infoMap['Number of Tables']
        tableEltsTxt=tableEltsTxt.replace('	','')
        tableEltsTxt=tableEltsTxt.replace('\n','')
        tableElts[0].text=tableEltsTxt
    submitdate=swiftxmlparser.xpath('//Submitted')
    if 'Original Received Date' in infoMap:
        submitdateTxt=infoMap['Original Received Date']
        submitdateTxt=submitdateTxt.replace('	','')
        submitdateTxt=submitdateTxt.replace('\n','')
        splittedDates1=submitdateTxt.split(' ')
        newAcceptedDate1=splittedDates1[2]+'-'+months[splittedDates1[1]]+'-'+days[splittedDates1[0]]
        submitdate[0].text=newAcceptedDate1
    rrdate=swiftxmlparser.xpath('//Revised')
    if 'Revision Received Date' in infoMap:
        rrdateTxt=infoMap['Revision Received Date']
        rrdateTxt=rrdateTxt.replace('	','')
        rrdateTxt=rrdateTxt.replace('\n','')
        splitdate1=rrdateTxt.split(' ')
        if(len(splitdate1)>3):
            newAcceptDate1=splitdate1[3]+'-'+months[splitdate1[2]]+'-'+days[splitdate1[1]]
        else:
            newAcceptDate1=splitdate1[2]+'-'+months[splitdate1[1]]+'-'+days[splitdate1[0]]
        rrdate[0].text=newAcceptDate1
    acceptDate=swiftxmlparser.xpath('//EditorialDecision')
    if 'Accepted Date' in infoMap:
        acceptDateTxt=infoMap['Accepted Date']
        acceptDateTxt=acceptDateTxt.replace('	','')
        acceptDateTxt=acceptDateTxt.replace('\n','')
        splittedDates=acceptDateTxt.split(' ')
        newAcceptedDate=splittedDates[2]+'-'+months[splittedDates[1]]+'-'+days[splittedDates[0]]
        acceptDate[0].text=newAcceptedDate
    wordcount=swiftxmlparser.xpath('//WordCount')
    wordcountTxt=etree.Element('WordCount')
    wordcountTxt.text='100'
    wordcount.append(wordcountTxt)
    emails=swiftxmlparser.xpath('//Email[@Primary="true"]')
    if 'Email' in infoMap:
        emailTxt=infoMap['Email']
        emailTxt=emailTxt.replace('	','')
        emailTxt=emailTxt.replace('\n','')
        emails[0].text=emailTxt
    phone=swiftxmlparser.xpath('//Telephone[@Primary="true"]')
    if 'Phone' in infoMap:
        phoneTxt=infoMap['Phone']
        phoneTxt=phoneTxt.replace('	','')
        phoneTxt=phoneTxt.replace('\n','')
        phone[0].text=phoneTxt
    ins=swiftxmlparser.xpath('//Institution')
    if 'Institution' in infoMap:
        insTxt=infoMap['Institution']
        insTxt=insTxt.replace('	','')
        insTxt=insTxt.replace('\n','')
        for i in ins:
            i.text=insTxt
    zipcode=swiftxmlparser.xpath('//PostCode')
    if 'Zip/Postal Code' in infoMap:
        zipcodeTxt=infoMap['Zip/Postal Code']
        zipcodeTxt=zipcodeTxt.replace(' ','')
        zipcodeTxt=zipcodeTxt.replace('\n','')
        for z in zipcode:
            z.text=zipcodeTxt
    dept=swiftxmlparser.xpath('//Department')
    address=swiftxmlparser.xpath('//Address1')
    town=swiftxmlparser.xpath('//Town')
    county=swiftxmlparser.xpath('//County')
    country=swiftxmlparser.xpath('//Country')
    sup_note=swiftxmlparser.xpath('//SupplementaryData')
    note_1='Note  1.'
    for note in infoMap:
        if isinstance(note, str) and re.search(note_1, note):
            swiftxmlparser.find('//SupplementaryData').text='true'
            xmlPath=zipFile.split('.')[0]
            swiftxmlparser.write(xmlPath+'.xml')
            break
        else:
            swiftxmlparser.find('//SupplementaryData').text='false'
            xmlPath=zipFile.split('.')[0]
            swiftxmlparser.write(xmlPath+'.xml')
    if 'Address' in infoMap:
        try:
            addressTxt=infoMap['Address']
            addressTxt=addressTxt.replace(' ','')
            addressTxt=addressTxt.replace('\n','')
            splittedAdd=addressTxt.split(',')
            dept[0].text=splittedAdd[0]
            address[0].text=splittedAdd[1]
            town[0].text=splittedAdd[2]
            county[0].text=splittedAdd[3]
            country[0].text=splittedAdd[4]
        except Exception as ex:
            print (ex)
    #Remove Contact ORCID value
    contacts=swiftxmlparser.xpath('//Contact')
    for c in contacts:
        try:
            c.attrib['ORCID']=''
        except Exception as exxx:
            print (exxx)
    gname=metaxmlparser.xpath('//contrib[not(@corresp="yes")]/name/given-names')
    sname=metaxmlparser.xpath('//contrib[not(@corresp="yes")]/name/surname')
    emailIDs=metaxmlparser.xpath('//contrib[not(@corresp="yes")]/email')
    surnames=swiftxmlparser.xpath('//SurName')
    foreNames=swiftxmlparser.xpath('//ForeName')
    gname1=metaxmlparser.xpath('//contrib[@corresp="yes"]/name/given-names')
    sname1=metaxmlparser.xpath('//contrib[@corresp="yes"]/name/surname')
    emailIDs1=metaxmlparser.xpath('//contrib[(@corresp="yes")]/email')
    affSurname=metaxmlparser.xpath('//contrib/name/surname')
    affMap={}
    authorAffForInstitution={}
    for a in affSurname:
        try:
            name=a.getparent()
            if(name.tag=='name'):
                xrefElt=name.getnext()
                affMap[a.text]=xrefElt.attrib['rid']
        except Exception as eee:
            print (eee)
    for m in affMap:
        try:
            singleAffs=metaxmlparser.xpath('//aff[@id="'+affMap[m]+'"]')
            comStr=stringify_children(singleAffs[0])
            comStr=comStr.replace('\n','')
            comStr=comStr.replace('\t','')
            comStr=comStr.replace(';','')
            comStr=comStr.replace(',','')
            comStr=comStr.lstrip()
            if(comStr[0].isdigit()==True):
                comStr=comStr[2:]
            authorAffForInstitution[m]=comStr
        except Exception as exex:
            print (exex)
            

    orcidValues={}
    '''
    for index,f in enumerate(foreNames):
        author=f.text+' '+surNames[index].text
        authors.append(author)
    '''
    #readtxt('//j-fs01/OUP_Journals-L/OUP_File_Autodownload/jexbot-308412/doc/'+baseFileName)          

    for index,f in enumerate(foreNames):
        try:
            if index!=0:
                f.text=gname[index-1].text
                f.getnext().text=sname[index-1].text
                f.getparent().attrib['orginal']='yes'
                try:
                    urlElts=gname[index-1].getparent().getparent().xpath('.//uri[@content-type="orcid"]')
                    urlTxt=urlElts[0].text
                    orcids=urlTxt.split('/')
                    orcid=orcids[len(orcids)-1]
                    '''
                    if 'Corresponding' in f.getparent().attrib:
                        f.getparent().attrib['ORCID']=orcid
                    '''
                    orcidValues[gname[index-1].text+'~'+sname[index-1].text]=orcid
                except Exception as exx:
                    print (exx)
                
                try:
                    emails=f.getparent().xpath('./Emails/Email')
                    emails[0].text=emailIDs[index-1].text
                except Exception as ee:
                    print (ee)
            try:
                emails=f.getparent().xpath('./Emails/Email')
                emails[0].text=emailIDs1[index].text
            except Exception as ex4:
                print(ex4)
        except Exception as ex:
            print (ex)
    firstForeName=foreNames[0]
    firstsurName=surnames[0]
    try:
        firstForeName.text=gname1[0].text 
        firstsurName.text=sname1[0].text
        firstForeName.getparent().attrib['orginal']='yes'
        urlElts1=gname1[0].getparent().getparent().xpath('./uri[@content-type="orcid"]')
        if(len(urlElts1)>0):
            urlTxt1=urlElts1[0].text
            orcids1=urlTxt1.split('/')
            orcid1=orcids1[len(orcids1)-1]
            orcidValues[gname1[0].text+'~'+sname1[0].text]=orcid1
    except Exception as ex6:
        print (ex6)
    try:
        for index,f1 in enumerate(foreNames):
            f1parent=f1.getparent()
            if 'ORCID' in f1parent.attrib and f1.text+'~'+f1.getnext().text in orcidValues:
                f1parent.attrib['ORCID']=orcidValues[f1.text+'~'+f1.getnext().text]
    except Exception as ex3:
        print(ex3)     
    contactsRem=swiftxmlparser.xpath('//Contact[not(@orginal)]')
    for r in contactsRem:
        r.getparent().remove(r)
    contactsRemAtt=swiftxmlparser.xpath('//Contact[@orginal]')
    for att in contactsRemAtt:
        att.attrib.pop('duplicate',None)
    
    
    baseFileName=os.path.basename(zipFile)
    baseFileName=baseFileName.replace('.zip','')
    docxpath=zipFile.replace('.zip','/doc')
    try:
        dir_list = os.listdir(docxpath)
    except Exception as ex7:
        print(ex7)
    docxpath1=''
    try:
        for d in dir_list:
            if '.docx' in d:
                docxpath1=d
                break
    except Exception as ex8:
        print(ex8)
   # readtxt(docxpath+'/'+docxpath1,swiftxmlparser)
    
    swiftSurname=swiftxmlparser.xpath('//SurName')
    for s in swiftSurname:
        try:
            #print (s.text)
            instStr=authorAffForInstitution[s.text]
            instStr=instStr.replace('\n','')
            instStr=instStr.replace('\t','')
            parentElt=s.getparent()
            instElt=parentElt.xpath('./Addresses/Address/Institution')
            instElt[0].text=instStr
        except Exception as ee:
            print (ee)
            
    contactsElts=swiftxmlparser.xpath('//Contact[@Corresponding="false"]')
    
    for c in contactsElts:
        try:
            address=c.xpath('.//Address1')
            for a in address:
                a.text=''
            department=c.xpath('.//Department')
            for d in department:
                d.text=''
            town=c.xpath('.//Town')
            for t in town:
                t.text=''
            county=c.xpath('.//County')
            for cou in county:
                cou.text=''
            country=c.xpath('.//Country')
            for co in country:
                co.text=''
            postalCode=c.xpath('.//PostCode')
            for p in postalCode:
                p.text=''
        except Exception as ee:
            print (ee)
    
    contactsElts1=swiftxmlparser.xpath('//Contact')
    
    for e in contactsElts1:
        try:
            e.attrib.pop('orginal',None)
        except Exception as eeee:
            print(eeee)
    level_from_style_name = {f'Heading {i}': i for i in range(1)}
    def format_levels(cur_lev):
        levs = [str(l) for l in cur_lev if l != 0]
        return '.'.join(levs)
    try:
        d = docx.Document(doc_Path)
        move_path=newPath+'/pdf'
        current_levels = [0] * 5
        full_text = []

        for p in d.paragraphs:
            if p.style.name not in level_from_style_name:
                full_text.append(p.text)
            else:
                level = level_from_style_name[p.style.name]
                current_levels[level] += 1
                for l in range(level + 1, 10):
                    current_levels[l] = 0
                full_text.append(format_levels(current_levels) + ' ' + p.text)
        if 'Manuscript Instructions Template (for journals without integrated submission systems)\n'==full_text[0]:
            shutil.move(doc_Path,move_path)
            dir_list_mov=os.listdir(move_path)
            for wordfilePath in dir_list_mov:
                if '.docx' in wordfilePath:
                    mov_word_file=wordfilePath
            word_FilePath=move_path+'/'+mov_word_file
            field_val=docx2txt.process(word_FilePath)
            list_of=field_val.split('\n\n')
            nxt_word=list_of[list_of.index('Member discount (%):')+1]
            nxt_word1=list_of[list_of.index('Colour discount (%):')+1]
            nxt_word2=list_of[list_of.index('Page discount (%): ')+1]
            nxt_word3=list_of[list_of.index('Special Issue?')+1]
            nxt_word4=list_of[list_of.index('Supplement Issue?')+1]
            nxt_word5=list_of[list_of.index('Embargo?')+1]
            nxt_word6=list_of[list_of.index('Press release?')+1]
            nxt_word7=list_of[list_of.index('Free to view?')+1]
            nxt_word8=list_of[list_of.index("Editor's Choice?")+1]
            nxt_word9=list_of[list_of.index('Linked Paper?')+1]
            nxt_word10=list_of[list_of.index("Skip ‘Advance Access’ for accepted manuscript?")+1]
            nxt_word11=list_of[list_of.index("Skip 'Advance Access' for final, typeset manuscript?")+1]
            nxt_word12=list_of[list_of.index('Supplement licence: CC BY (no author charge)?')+1]
            nxt_word13=list_of[list_of.index('Supplement licence: CC BY-NC (no author charge)?')+1]
            try:
                oaDiscount=swiftxmlparser.xpath('//Instructions/OA_Discount')
                if oaDiscount is not None:
                    oaDiscount[0].text=nxt_word
                if nxt_word is None:
                    oaDiscount[0].text='0'
            except Exception as ex0:
                print(ex0)
            try:
                colour_Discount=swiftxmlparser.xpath('//Instructions/Colour_Discount')
                if colour_Discount is not None:
                    colour_Discount[0].text=nxt_word1
            except Exception as ex1:
                print(ex1)
            try:
                pages_Waived=swiftxmlparser.xpath('//Instructions/Pages_Waived')
                if pages_Waived is not None:
                    pages_Waived[0].text=nxt_word2
            except Exception as ex2:
                 print(ex2)
            try:
                special_Issue=swiftxmlparser.xpath('//Instructions/Special_Issue')
                if special_Issue is not None:
                
                    if 'No'==nxt_word3:
                        special_Issue[0].text='0'
                    if 'Yes'==nxt_word3:
                        special_Issue[0].text='1'
            except Exception as ex3:
                     print(ex3)
            try:
                supplement=swiftxmlparser.xpath('//Instructions/Supplement')
                if supplement is not None:
                    if 'No'==nxt_word4:
                        supplement[0].text='0'
                    if 'Yes'==nxt_word4:
                        supplement[0].text='1'
            except Exception as ex4:
                     print(ex4)
            try:
                embargo=swiftxmlparser.xpath('//Instructions/Embargo')
                if embargo is not None:
                    if 'No'==nxt_word5:
                        embargo[0].text='0'
                        #embargo[0].text='False'
                    if 'Yes'==nxt_word5:
                        embargo[0].text='1'
                        #embargo[0].text='True'
            except Exception as ex5:
                     print(ex5)
            pressRelease=swiftxmlparser.xpath('//Instructions/PressRelease')
            if pressRelease is not None:
                if 'No'==nxt_word6:
                    pressRelease[0].text='0'
                    #pressRelease[0].text='False'
                if 'Yes'==nxt_word6:
                    pressRelease[0].text='1'
                    #pressRelease[0].text='True'
            freeToView=swiftxmlparser.xpath('//Instructions/FreeToView')
            if freeToView is not None:
                if 'No'==nxt_word7:
                    freeToView[0].text='0'
                    #freeToView[0].text='False'
                if 'Yes'==nxt_word7:
                    freeToView[0].text='1'
                    #freeToView[0].text='True'
            editorChoice=swiftxmlparser.xpath('//Instructions/EditorChoice')
            if editorChoice is not None:
                if 'No'==nxt_word8:
                    editorChoice[0].text='0'
                    #editorChoice[0].text='False'
                if 'Yes'==nxt_word8:
                    editorChoice[0].text='1'
                    #editorChoice[0].text='True'
            linkedPaper=swiftxmlparser.xpath('//Instructions/LinkedPaper')
            if linkedPaper is not None:
                if 'No'==nxt_word9:
                    linkedPaper[0].text='0'
                    #linkedPaper[0].text='False'
                if 'Yes'==nxt_word9:
                    linkedPaper[0].text='1'
                    #linkedPaper[0].text='True'
            skipUnCorrManStep=swiftxmlparser.xpath('//Instructions/SkipUnCorrManStep')
            if skipUnCorrManStep is not None:
                if 'No'==nxt_word10:
                    skipUnCorrManStep[0].text='0'
                    #skipUnCorrManStep[0].text='False'
                if 'Yes'==nxt_word10:
                    skipUnCorrManStep[0].text='1'
                    #skipUnCorrManStep[0].text='True'
            skipAdvncAccessStep=swiftxmlparser.xpath('//Instructions/SkipAdvncAccessStep')
            if skipAdvncAccessStep is not None:
                if 'No'==nxt_word11:
                    skipAdvncAccessStep[0].text='0'
                    #skipAdvncAccessStep[0].text='False'
                if 'Yes'==nxt_word11:
                    skipAdvncAccessStep[0].text='1'
                    #skipAdvncAccessStep[0].text='True'
            cC_BY_No_Charge=swiftxmlparser.xpath('//Instructions/CC_BY_No_Charge')
            if cC_BY_No_Charge is not None:
                if 'No'==nxt_word12:
                    cC_BY_No_Charge[0].text='0'
                    #cC_BY_No_Charge[0].text='False'
                if 'Yes'==nxt_word12:
                    cC_BY_No_Charge[0].text='1'
                    #cC_BY_No_Charge[0].text='True'
            cC_BY_NC_No_Charge=swiftxmlparser.xpath('//Instructions/CC_BY_NC_No_Charge')
            if cC_BY_NC_No_Charge is not None:
                if 'No'==nxt_word13:
                    cC_BY_NC_No_Charge[0].text='0'
                    #cC_BY_NC_No_Charge[0].text='False'
                if 'Yes'==nxt_word13:
                    cC_BY_NC_No_Charge[0].text='1'
                    #cC_BY_NC_No_Charge[0].text='True'
    except Exception as ex5:
        print(ex5)
    notesElts=swiftxmlparser.xpath('//Note[@Code="TYPESETTER"]')
    notesElts[0].text=notes
    swiftXMLFile=os.path.basename(zipFile)
    swiftXMLFile=swiftXMLFile.replace('.zip','.xml')
    swiftxmlparser.write(createdSwiftXML+''+swiftXMLFile,xml_declaration=True, encoding='utf-8',pretty_print=True)
    fornamesAuth=swiftxmlparser.xpath('//Contact/ForeName')
    for authname in fornamesAuth:
        aTxt=authname.text
        affTxt=metaxmlparser.xpath(f"//name[given-names[text()='{aTxt}']]/following-sibling::xref/@rid")
        for affno in affTxt:
            if affno.startswith('aff'):
                ringGold=metaxmlparser.xpath(f"//aff[@id='{affno}']//institution-id[@institution-id-type='Ringgold']/text()")
                if ringGold:
                    ringGold=ringGold[0]
                    ringgold_id=swiftxmlparser.xpath(f"//Contact[ForeName[text()='{aTxt}']]//RinggoldID")
                    if ringgold_id:
                        ringgold_id[0].text = ringGold
                        swiftxmlparser.write(xmlPath+'.xml')
    ins=swiftxmlparser.xpath('//Addresses/Address/Institution')
    for i in ins:
        inTxt=i.text
        iTxt=inTxt.rstrip()
        i.text=iTxt
    swiftxmlparser.write(xmlPath+'.xml')
    xmlStr=''
    with open(createdSwiftXML+''+swiftXMLFile,'r',encoding="utf-8") as f:
        xmlStr=f.read()
        xmlStr=xmlStr.replace('xmlns:temp="','xmlns="')
        with open(createdSwiftXML+''+swiftXMLFile,'w',encoding="utf-8") as w:
            w.write(xmlStr)
        f.close()
        # Find the sibling element
        '''parent = target_element.getparent()
        siblings = parent.getchildren()
        index = siblings.index(target_element)
        target_element = root.xpath('//target_element[text()="Specific Text"]')[0]  # Replace with your tag and text
        print(f"Target element found with text: {target_element.text}")

        if index + 1 < len(siblings):  # Check if a neighbor exists
            neighbor_element = siblings[index + 1]
            print(f"Original neighbor element text: {neighbor_element.text}")'''

    swiftXML=createdSwiftXML+''+swiftXMLFile
    try:
        output = subprocess.run(["C:/Program Files/Adobe/Adobe InDesign 2020/Plug-Ins/AutoComp/UTF8.exe",createdSwiftXML+'' +swiftXMLFile], capture_output=True)
        upload_xmlfile(swiftXML)
    except Exception as ex3:
        print (ex3)
    
    baseName=os.path.basename(swiftXML)
    zipFolderPath=os.path.dirname(swiftXML)
    fileName=baseName.replace('.xml','')
    zipFolderName=zipFolderPath+'/'+fileName
    #shutil.move(swiftXML,zipFolderName)
    zipName=zipFolderPath+'/Upload_zip_file/'+fileName+'.zip'
    with zipfile.ZipFile(zipName, 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipdir(zipFolderName, zipf)
        zipf.close()
    try: 
        myZipFiles=ReadZipFilesFromFTP()
        if(os.path.basename(zipName) not in myZipFiles):
            upload_file(zipName)
    except Exception as ex5:
        print(ex5)
        
def ReadZipFilesFromFTP():
    myZipFiles=[]
    try:
        s = ftplib.FTP('ftp.newgenimaging.com','oupjournals', 'uEmg3n9upJ0UrNl')
        s.cwd('/From_OUP/EXBOTJ/reprocess')
        files = s.nlst()
        for f in files:
            myZipFiles.append(f)
    except Exception as e:
        print (e)
    return myZipFiles
def ReadXmlFilesFromFTP():
    myXmlFiles=[]
    try:
        x = ftplib.FTP('ukftp-auto.oup.com','jautomaton','send2me')
        x.cwd('/TS.v1')
        xmlfiles = x.nlst()
        for xl in xmlfiles:
            myXmlFiles.append(xl)
    except Exception as exxxxx:
        print (exxxxx) 
    return myXmlFiles
def zipdir(path, ziph):
    # ziph is zipfile handle
    for root, dirs, files in os.walk(path):
        for file in files:
            ziph.write(os.path.join(root, file), 
                       os.path.relpath(os.path.join(root, file), 
                                       os.path.join(path, '..')))

def Ftpconnection(myMap): 
    ftp = ftplib.FTP("ftp.newgenimaging.com",timeout=10)
    ftp.login("oupjournals", "uEmg3n9upJ0UrNl")
    
    
    file_names = sorted(ftp.nlst('From_OUP/EXBOTJ/Accepted_Manuscripts/'), key=lambda x: ftp.voidcmd(f"MDTM {x}"))
    intCnt=0
    downloadArr=[]
    for file_name in reversed(file_names):
        if intCnt>=5:
            break
        baseFileName=os.path.basename(file_name)
        if baseFileName.startswith('jexbot-'):
            print(baseFileName)
            downloadZip='//j-fs01/OUP_Journals-L/OUP_File_Autodownload/'+baseFileName
            if downloadZip not in myMap:
                urllib.request.urlretrieve('ftp://oupjournals:uEmg3n9upJ0UrNl@ftp.newgenimaging.com/'+file_name, downloadZip) 
                intCnt=intCnt+1
                downloadArr.append(downloadZip)
    return downloadArr
def show_success(message):
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo("Success", f"{message} Manuscript inserted successfully!")  # Show success popup
    root.destroy()
def show_warning(message):
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo("Warning", message)
    root.destroy()
count=0
ftpName=''
userName=''
password=''
myMap={}
manuscriptnames=[]
upl_file='//j-fs01/OUP_Journals-L/OUP_File_Autodownload/Upload_zip_file/Upload_xml_details.txt'
while True:
    try:
        count=count+1
        '''if count==6:
            warn='This Manuscripts not inserted '+manuscriptnames
            show_warning(warn,"ERROR")
            sys.exit()'''
        if count%5==0:
            downloadZipFiles=Ftpconnection(myMap)
            for downloadZipFile in downloadZipFiles:
                if(os.path.exists(downloadZipFile)==True):
                    if downloadZipFile not in myMap:
                        myMap[downloadZipFile]=downloadZipFile
                        MainFunction(downloadZipFile)        
        time.sleep(15)
    except Exception as exe:
        print(exe)